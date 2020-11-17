import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import date
today = date.today()
today=today.strftime('%m/%d/%Y, %H:%M:%S')
doc =  Document()
doc.add_heading('Hospital',0)
thanks =doc.add_paragraph()
thanks.add_run('Your Health Our Priority \n').bold = True
doc_para = doc.add_paragraph()
doc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc_para.add_run('Patient ID: anubhavgupta2260@gmail.com	AGE:19	GENDER:M \n').font.size = Pt(16)
doc_para.add_run('Prescription : 		').font.size = Pt(14)
doc_para.add_run('\n\n\n\n\n\n\n\n\n').font.size = Pt(14)
#doc_para.add_run(txn.rjust(16)).font.size = Pt(14)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('Dermatologist  \n ')
footer.add_run(' Doctor email : SamwellTarly@gmail.com   \n')
#footer.add_run(' Ph: 99#######,91########\n')
doc.save('printreceipt.docx')
os.startfile('printreceipt.docx', 'print')
